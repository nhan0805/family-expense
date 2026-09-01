from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from PIL import Image

ROOT = Path('/Users/hanhuynh/Downloads/Expense')
OUT = ROOT / 'Huong-dan-su-dung-Family-Expense.docx'
ASSETS = ROOT / 'docs-assets'

NAVY = '173E32'
GREEN = '137050'
LIGHT_GREEN = 'E8F3ED'
PALE = 'F4F7F5'
GRAY = '667085'
LIGHT_GRAY = 'E5E7EB'
WHITE = 'FFFFFF'
RED = 'B42318'


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{margin}'))
        if node is None:
            node = OxmlElement(f'w:{margin}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    cant_split.set(qn('w:val'), 'true')
    tr_pr.append(cant_split)


def set_table_widths(table, widths):
    table.autofit = False
    # Đánh dấu hàng đầu để trình đọc màn hình nhận diện cấu trúc bảng.
    # Với các callout một hàng, thuộc tính này vẫn an toàn và không đổi bố cục.
    if table.rows:
        set_repeat_table_header(table.rows[0])
    tbl_grid = table._tbl.tblGrid
    grid_cols = tbl_grid.findall(qn('w:gridCol'))
    for idx, width in enumerate(widths):
        if idx < len(grid_cols):
            grid_col = grid_cols[idx]
        else:
            grid_col = OxmlElement('w:gridCol')
            tbl_grid.append(grid_col)
        grid_col.set(qn('w:w'), str(round(width * 1440)))
    for row in table.rows:
        prevent_row_split(row)
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(round(width * 1440)))
            tc_w.set(qn('w:type'), 'dxa')
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(round(sum(widths) * 1440)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '180' if len(widths) == 1 else '120')
    tbl_ind.set(qn('w:type'), 'dxa')


def set_font(run, size=None, color=None, bold=None, italic=None, name='Calibri'):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), name)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Trang ')
    set_font(run, size=9, color=GRAY)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    keep_with_next(p)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        set_font(p.add_run(text))
    return p


def add_bullet(doc, text, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(text, style=style)
    return p


def set_number_value(paragraph, value):
    numbering = paragraph.part.numbering_part.element
    nums = numbering.findall(qn('w:num'))
    new_num_id = max([int(n.get(qn('w:numId'))) for n in nums] + [0]) + 1
    base_num = next((n for n in nums if n.get(qn('w:numId')) == '5'), nums[0])
    abstract = base_num.find(qn('w:abstractNumId')).get(qn('w:val'))
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(new_num_id))
    abstract_ref = OxmlElement('w:abstractNumId')
    abstract_ref.set(qn('w:val'), abstract)
    num.append(abstract_ref)
    override = OxmlElement('w:lvlOverride')
    override.set(qn('w:ilvl'), '0')
    start = OxmlElement('w:startOverride')
    start.set(qn('w:val'), str(value))
    override.append(start)
    num.append(override)
    numbering.append(num)
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    num_id = OxmlElement('w:numId')
    num_id.set(qn('w:val'), str(new_num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id)


def add_step(doc, number, title, detail):
    p = doc.add_paragraph(style='List Number')
    set_number_value(p, number)
    r = p.add_run(title)
    set_font(r, bold=True, color=NAVY)
    r = p.add_run(f' — {detail}')
    set_font(r)
    return p


def add_callout(doc, title, text, kind='note'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, 'FFF4E5' if kind == 'warning' else LIGHT_GREEN)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, bold=True, color=RED if kind == 'warning' else GREEN)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.runs[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, path, caption, width=6.15):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set('descr', caption)
    picture._inline.docPr.set('title', caption)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    set_font(cap.runs[0], size=9, color=GRAY, italic=True)


def add_quick_table(doc):
    data = [
        ('Tổng quan', 'KPI thu/chi/ròng, biểu đồ, giao dịch dự kiến đến hạn.'),
        ('Giao dịch', 'Tìm kiếm, lọc, sắp xếp, xem, sửa, sao chép và xóa.'),
        ('Danh mục', 'Quản lý mục đích chi, loại chi phí, phương thức thanh toán.'),
        ('Thành viên', 'Quản lý gia đình và quyền truy cập của các thành viên.'),
        ('Dữ liệu', 'Tải template, import và xuất dữ liệu giao dịch bằng Excel.'),
        ('Nút +', 'Thêm giao dịch mới từ hầu hết màn hình.'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [1.55, 4.95])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(('Khu vực', 'Công dụng')):
        shade(hdr.cells[idx], NAVY)
        p = hdr.cells[idx].paragraphs[0]
        r = p.add_run(text)
        set_font(r, bold=True, color=WHITE)
    for label, detail in data:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        set_font(cells[0].paragraphs[0].runs[0], bold=True, color=NAVY)
        set_font(cells[1].paragraphs[0].runs[0])
    for row in table.rows:
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
login_crop = ASSETS / 'login-cropped.png'
with Image.open(ASSETS / 'login.png') as source:
    source.crop((170, 50, 470, 315)).save(login_crop)
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in [
    ('Heading 1', 16, GREEN, 18, 10),
    ('Heading 2', 13, GREEN, 14, 7),
    ('Heading 3', 12, NAVY, 10, 5),
]:
    st = styles[name]
    st.font.name = 'Calibri'
    st._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri')
    st._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ['List Bullet', 'List Bullet 2', 'List Number']:
    st = styles[list_name]
    st.font.name = 'Calibri'
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

header = section.header.paragraphs[0]
header.text = 'FAMILY EXPENSE  |  HƯỚNG DẪN NGƯỜI DÙNG'
set_font(header.runs[0], size=8.5, color=GRAY, bold=True)
footer = section.footer.paragraphs[0]
page_number(footer)

# Cover — editorial_cover pattern.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(84)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FAMILY EXPENSE')
set_font(r, size=11, color=GREEN, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
r = p.add_run('HƯỚNG DẪN SỬ DỤNG')
set_font(r, size=30, color=NAVY, bold=True)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(30)
r = p.add_run('Quản lý thu chi gia đình rõ ràng, nhanh chóng và có hỗ trợ AI')
set_font(r, size=14, color=GRAY)
add_callout(doc, 'Tài liệu dành cho người dùng cuối', 'Hướng dẫn này mô tả phiên bản ứng dụng cập nhật ngày 27/08/2026, áp dụng cho cả giao diện máy tính và điện thoại.')
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(120)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Phiên bản 1.1  •  Tháng 08/2026')
set_font(r, size=10, color=GRAY, italic=True)
doc.add_page_break()

add_heading(doc, 'Mục lục', 1)
toc_items = [
    '1. Bắt đầu sử dụng', '2. Làm quen với màn hình chính', '3. Tổng quan tài chính',
    '4. Danh sách giao dịch', '5. Thêm và chỉnh sửa giao dịch',
    '6. Nhập giao dịch bằng AI', '7. Giao dịch dự kiến và xác nhận đến hạn',
    '8. Quản lý danh mục', '9. Import và xuất dữ liệu Excel',
    '10. Gia đình và thành viên', '11. Cài ứng dụng trên điện thoại',
    '12. Quy tắc dữ liệu quan trọng', '13. Mẹo sử dụng và xử lý sự cố',
]
for item in toc_items:
    p = add_body(doc, item)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(3)
add_callout(doc, 'Bắt đầu nhanh', 'Đăng nhập → kiểm tra kỳ trên Tổng quan → bấm nút “+” → nhập nội dung trực tiếp hoặc nhấn icon AI để nhận gợi ý → kiểm tra thông tin và lưu.')

add_heading(doc, '1. Bắt đầu sử dụng', 1)
add_heading(doc, '1.1 Đăng nhập', 2)
add_step(doc, 1, 'Mở ứng dụng', 'truy cập địa chỉ Family Expense bằng trình duyệt trên máy tính hoặc điện thoại.')
add_step(doc, 2, 'Nhập Email và Mật khẩu', 'sử dụng tài khoản đã đăng ký, sau đó chọn “Tiếp tục”.')
add_step(doc, 3, 'Chờ tải dữ liệu', 'ứng dụng sẽ mở dữ liệu của gia đình mà tài khoản đang tham gia.')
add_figure(doc, login_crop, 'Hình 1. Màn hình đăng nhập Family Expense', width=3.5)
add_heading(doc, '1.2 Các cách truy cập tài khoản khác', 2)
add_bullet(doc, 'Đăng ký: tạo tài khoản mới bằng email và mật khẩu.')
add_bullet(doc, 'Quên mật khẩu: nhận liên kết đặt lại mật khẩu qua email.')
add_bullet(doc, 'Magic link: nhận liên kết đăng nhập qua email mà không cần nhập mật khẩu.')
add_bullet(doc, 'Đăng xuất: dùng nút Đăng xuất ở thanh trên; màn hình tạo gia đình cũng có nút Đăng xuất để đổi tài khoản.')
add_callout(doc, 'Lưu ý bảo mật', 'Không chia sẻ mật khẩu hoặc liên kết đăng nhập qua email. Khi dùng thiết bị công cộng, hãy đăng xuất sau khi hoàn tất.', 'warning')

add_heading(doc, '2. Làm quen với màn hình chính', 1)
add_body(doc, 'Ứng dụng có sáu khu vực chính. Trên máy tính, menu nằm bên trái; trên điện thoại, menu nằm ở thanh dưới cùng.')
add_quick_table(doc)
add_body(doc, 'Trên điện thoại, bấm biểu tượng menu nếu cần mở thanh điều hướng. Nút tròn “+” ở góc dưới dùng để thêm giao dịch mới nhanh chóng.')

add_heading(doc, '3. Tổng quan tài chính', 1)
add_heading(doc, '3.1 Chọn khoảng thời gian', 2)
add_body(doc, 'Ở đầu màn hình Tổng quan, chọn riêng Tháng và Năm. Dùng “Tháng trước/Tháng này” để chuyển nhanh; trên điện thoại các nút được rút gọn thành “Trước/Nay” và nằm cùng hàng với Tháng, Năm.')
add_heading(doc, '3.2 Đọc các KPI', 2)
add_bullet(doc, 'Tổng thu: tổng các giao dịch thu nhập thực tế trong kỳ.')
add_bullet(doc, 'Tổng chi: tổng các khoản chi thực tế trong kỳ.')
add_bullet(doc, 'Giá trị ròng: Tổng thu trừ Tổng chi.')
add_body(doc, 'Số tiền trong KPI được hiển thị đầy đủ. Hai biểu đồ có thể rút gọn nhãn tiền theo K hoặc M để tránh chồng chữ, nhưng giá trị dữ liệu không thay đổi.')
add_heading(doc, '3.3 Biểu đồ và danh sách gần đây', 2)
add_bullet(doc, 'Chi tiêu theo mục đích chi: biểu đồ tròn so sánh các mục đích trong kỳ.')
add_bullet(doc, 'Chi tiêu theo loại chi phí: biểu đồ cột so sánh các loại; trên điện thoại có thể cuộn ngang bên trong card.')
add_bullet(doc, 'Xu hướng 5 tháng: theo dõi biến động chi ròng trong năm tháng gần nhất.')
add_bullet(doc, 'Giao dịch thực tế trong tháng: xem nhanh các giao dịch đã phát sinh; chọn “Xem tất cả” để chuyển sang danh sách đã lọc theo tháng/năm.')

add_heading(doc, '4. Danh sách giao dịch', 1)
add_body(doc, 'Màn hình Giao dịch hiển thị ngày, nội dung, mục đích chi, loại chi phí, phương thức thanh toán và số tiền. Trạng thái được ẩn khỏi danh sách và bộ lọc để giao diện gọn hơn; dữ liệu trạng thái vẫn được giữ để Dashboard tính đúng.')
add_heading(doc, '4.1 Tìm kiếm, lọc và sắp xếp', 2)
add_body(doc, 'Card “Chi ròng theo bộ lọc” ở đầu màn hình luôn tính trên toàn bộ kết quả phù hợp, kể cả khi danh sách đang chỉ tải một phần để tối ưu tốc độ.')
add_callout(doc, 'Công thức chi ròng', 'Chi tiêu − Thu nhập. Vì vậy giao dịch thu nhập làm giảm tổng thay vì được cộng vào tổng chi.')
add_bullet(doc, 'Tìm kiếm theo nội dung hoặc ghi chú; có thể gõ không dấu, ví dụ “dien nuoc” vẫn tìm được “Điện nước”.')
add_bullet(doc, 'Sắp xếp theo ngày mới/cũ, số tiền cao/thấp hoặc nội dung A-Z.')
add_bullet(doc, 'Chọn “Bộ lọc chi tiết” để mở/thu gọn các trường loại giao dịch, mục đích chi, loại chi phí và phương thức thanh toán.')
add_bullet(doc, 'Lọc độc lập theo tháng và năm, hoặc nhập khoảng Từ ngày/Đến ngày.')
add_bullet(doc, 'Các chip cho biết điều kiện đang chọn; bấm chip để bỏ nhanh hoặc chọn “Xóa bộ lọc” để đặt lại toàn bộ.')
add_heading(doc, '4.2 Thao tác trên một giao dịch', 2)
add_bullet(doc, 'Chọn nội dung giao dịch để mở màn hình xem/chỉnh sửa.')
add_bullet(doc, 'Biểu tượng sao chép tạo một bản sao để nhập nhanh giao dịch tương tự.')
add_bullet(doc, 'Biểu tượng thùng rác xóa giao dịch sau khi bạn xác nhận.')
add_bullet(doc, 'Chủ gia đình có thể xóa mọi giao dịch; thành viên chỉ được xóa giao dịch do chính mình tạo.')
add_callout(doc, 'Mẹo', 'Khi cần tìm một khoản cũ, hãy kết hợp Năm + Tháng + một từ khóa ngắn trong nội dung. Cách này thường nhanh hơn nhập khoảng ngày.')

add_heading(doc, '5. Thêm và chỉnh sửa giao dịch', 1)
add_heading(doc, '5.1 Form giao dịch hợp nhất', 2)
add_step(doc, 1, 'Mở biểu mẫu', 'bấm nút “+” hoặc nút thêm giao dịch trên màn hình Giao dịch.')
add_step(doc, 2, 'Nhập nội dung', 'điền trực tiếp hoặc viết một câu tiếng Việt rồi nhấn icon AI để hệ thống gợi ý các trường.')
add_step(doc, 3, 'Hoàn thiện các trường', 'kiểm tra ngày, số tiền, loại giao dịch và ba trường phân loại bắt buộc.')
add_step(doc, 4, 'Mở tùy chọn nâng cao khi cần', 'trạng thái được tự chọn theo ngày; chỉ đổi khi giao dịch thực tế khác quy tắc ngày.')
add_step(doc, 5, 'Lưu giao dịch', 'nếu hệ thống phát hiện giao dịch tương tự, hãy kiểm tra cảnh báo trùng trước khi tiếp tục.')
add_heading(doc, '5.2 Ý nghĩa các trường', 2)
fields = [
    ('Ngày *', 'Ngày phát sinh hoặc dự kiến phát sinh.'),
    ('Số tiền (VND) *', 'Chỉ nhập số; ứng dụng tự định dạng dấu phân cách hàng nghìn.'),
    ('Loại giao dịch *', 'Chi tiêu hoặc Thu nhập.'),
    ('Nội dung *', 'Mô tả ngắn để nhập tay/tìm kiếm, hoặc câu tiếng Việt dùng cho nút gợi ý AI.'),
    ('Phương thức thanh toán *', 'Mặc định là Chuyển khoản; có thể chọn phương thức khác.'),
    ('Mục đích chi *', 'Nhóm mục tiêu sử dụng tiền.'),
    ('Loại chi phí *', 'Nhóm bản chất khoản chi.'),
    ('Trạng thái', 'Nằm trong Tùy chọn nâng cao; tự chọn Dự kiến cho ngày tương lai, ngược lại là Thực tế.'),
    ('Ghi chú', 'Nằm trong Tùy chọn nâng cao; thông tin bổ sung, không bắt buộc.'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
set_table_widths(table, [2.0, 4.5])
for i, text in enumerate(('Trường', 'Cách sử dụng')):
    shade(table.rows[0].cells[i], NAVY)
    set_font(table.rows[0].cells[i].paragraphs[0].add_run(text), bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
for label, detail in fields:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = detail
    set_font(cells[0].paragraphs[0].runs[0], bold=True, color=NAVY)
    set_font(cells[1].paragraphs[0].runs[0])
for row in table.rows:
    prevent_row_split(row)
    for cell in row.cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
add_callout(doc, 'Dấu sao (*)', 'Trường có dấu sao là bắt buộc. Nếu bỏ trống, ứng dụng sẽ hiển thị lỗi và chưa lưu giao dịch.')

add_heading(doc, '6. Nhập giao dịch bằng AI', 1)
add_body(doc, 'AI được tích hợp ngay trong form, không còn tab hoặc màn hình nhập riêng. Bạn vẫn có thể bỏ qua AI và nhập toàn bộ trường thủ công.')
add_step(doc, 1, 'Nhập một câu vào ô Nội dung', 'nêu ngày, số tiền, người/mục đích và phương thức thanh toán nếu có.')
add_step(doc, 2, 'Nhấn icon AI/Gợi ý AI', 'đợi hệ thống phân tích và điền đề xuất vào chính form hiện tại.')
add_step(doc, 3, 'Đọc cảnh báo', 'kiểm tra confidence, thông tin thiếu và mọi trường AI vừa đề xuất.')
add_step(doc, 4, 'Chỉnh sửa', 'đặc biệt là số tiền, ngày, mục đích chi, loại chi phí và phương thức thanh toán.')
add_step(doc, 5, 'Xác nhận và lưu', 'AI không tự lưu; database chỉ nhận giao dịch sau khi bạn nhấn nút lưu.')
add_heading(doc, '6.1 Ví dụ câu nhập tốt', 2)
add_bullet(doc, '“Hôm nay mua sữa cho Haku 450 nghìn, chuyển khoản.”')
add_bullet(doc, '“Ngày 30/8 nhận lương 25 triệu qua chuyển khoản.”')
add_bullet(doc, '“Tuần sau đóng học phí 3 triệu bằng tiền mặt.”')
add_bullet(doc, '“Được hoàn 250 nghìn tiền mua thuốc hôm qua.”')
add_heading(doc, '6.2 Cách giúp AI phân tích nhanh và chính xác', 2)
add_bullet(doc, 'Mỗi lần chỉ mô tả một giao dịch.')
add_bullet(doc, 'Ghi số tiền và đơn vị rõ ràng: 450 nghìn, 1,2 triệu hoặc 1.200.000.')
add_bullet(doc, 'Dùng tên danh mục gần với danh mục đang có trong ứng dụng.')
add_bullet(doc, 'Tránh câu quá dài, nhiều mệnh đề hoặc chứa nhiều khoản tiền khác nhau.')
add_callout(doc, 'Luôn kiểm tra trước khi lưu', 'AI có thể hiểu sai câu mơ hồ. Nếu thiếu số tiền, hệ thống sẽ để trống và cảnh báo; nếu thiếu ngày, AI có thể dùng ngày hiện tại kèm cảnh báo.', 'warning')

add_heading(doc, '7. Giao dịch dự kiến và xác nhận đến hạn', 1)
add_heading(doc, '7.1 Trạng thái được xác định thế nào?', 2)
add_bullet(doc, 'Ngày tương lai → Dự kiến.')
add_bullet(doc, 'Ngày hiện tại hoặc quá khứ → Thực tế.')
add_body(doc, 'Quy tắc này áp dụng khi nhập mới và giúp tránh ghi nhận sớm một khoản chưa phát sinh. Trạng thái vẫn được lưu trong dữ liệu dù không hiển thị dưới từng dòng trong danh sách.')
add_heading(doc, '7.2 Khi giao dịch dự kiến đến hạn', 2)
add_step(doc, 1, 'Mở Tổng quan', 'xem khu vực “Giao dịch dự kiến đến hạn”.')
add_step(doc, 2, 'Kiểm tra giao dịch', 'đảm bảo khoản tiền thực sự đã phát sinh.')
add_step(doc, 3, 'Bấm “Xác nhận thực tế”', 'đồng ý trong hộp thoại xác nhận.')
add_step(doc, 4, 'Hoàn tất', 'trạng thái được chuyển từ Dự kiến sang Thực tế và KPI/biểu đồ sẽ tính giao dịch theo kỳ phù hợp.')
add_callout(doc, 'Nếu khoản chưa phát sinh', 'Không xác nhận. Bạn có thể mở giao dịch để sửa ngày hoặc thông tin, hoặc giữ trạng thái Dự kiến để xử lý sau.')

add_heading(doc, '8. Quản lý danh mục', 1)
add_body(doc, 'Màn hình Danh mục gồm Mục đích chi, Loại chi phí và Phương thức thanh toán.')
add_heading(doc, '8.1 Quyền của người dùng', 2)
add_bullet(doc, 'Chủ gia đình: có thể thêm, đổi tên và xóa danh mục chưa được sử dụng.')
add_bullet(doc, 'Thành viên: chỉ có thể xem danh mục.')
add_heading(doc, '8.2 Thêm, sửa và xóa', 2)
add_bullet(doc, 'Thêm: chọn “Thêm” tại đúng nhóm, nhập tên và chọn “Lưu danh mục”.')
add_bullet(doc, 'Sửa: chọn biểu tượng bút chì, đổi tên và chọn “Lưu tên mới”.')
add_bullet(doc, 'Xóa: chọn biểu tượng thùng rác và xác nhận. Danh mục đã được giao dịch sử dụng sẽ không thể xóa.')
add_callout(doc, 'Khuyến nghị', 'Dùng tên danh mục ngắn, nhất quán và không tạo nhiều danh mục gần giống nhau. Điều này giúp bộ lọc và AI phân loại chính xác hơn.')

add_heading(doc, '9. Import và xuất dữ liệu Excel', 1)
add_heading(doc, '9.1 Tải và điền template', 2)
add_step(doc, 1, 'Mở “Dữ liệu”', 'chọn card “Tải template”.')
add_step(doc, 2, 'Mở file .xlsx', 'điền dữ liệu trong sheet “Giao dịch”; không đổi tên sheet hoặc tiêu đề cột.')
add_step(doc, 3, 'Dùng danh sách chọn', 'chọn loại giao dịch, trạng thái và danh mục từ dropdown có sẵn để tránh sai tên.')
add_body(doc, 'Template có các trường giống biểu mẫu tạo giao dịch, hướng dẫn điền và validation dữ liệu. Khi danh mục thay đổi, nên tải template mới để có danh sách cập nhật.')
add_heading(doc, '9.2 Kiểm tra và import', 2)
add_step(doc, 1, 'Chọn file', 'kéo-thả hoặc chọn file .xlsx đúng template Family Expense.')
add_step(doc, 2, 'Đọc kết quả kiểm tra', 'xem số dòng hợp lệ, có thể trùng và bị lỗi; kiểm tra phần preview trước khi ghi.')
add_step(doc, 3, 'Xử lý dòng lỗi', 'sửa file theo thông báo rồi chọn lại. File sai sheet, sai tiêu đề hoặc hỏng sẽ bị từ chối.')
add_step(doc, 4, 'Xác nhận import', 'ứng dụng chỉ gửi các dòng hợp lệ; mặc định bỏ qua dòng có thể trùng.')
add_callout(doc, 'Giới hạn template', 'Mỗi file hỗ trợ tối đa 1.000 dòng. Không đổi tên sheet “Giao dịch” hoặc tiêu đề cột.', 'warning')
add_heading(doc, '9.3 Xuất dữ liệu', 2)
add_step(doc, 1, 'Chọn “Xuất dữ liệu”', 'ứng dụng tạo file chứa toàn bộ giao dịch chưa xóa.')
add_step(doc, 2, 'Lưu file', 'chọn vị trí an toàn trên thiết bị.')
add_body(doc, 'File xuất gồm thông tin phân loại, phương thức thanh toán, ghi chú, nguồn nhập và dữ liệu audit liên quan. Có thể dùng để lưu trữ, đối soát hoặc phân tích thêm.')
add_callout(doc, 'Dữ liệu tài chính', 'File Excel có thể chứa thông tin nhạy cảm. Chỉ chia sẻ cho người được phép và tránh lưu trên thiết bị công cộng.', 'warning')

add_heading(doc, '10. Gia đình và thành viên', 1)
add_heading(doc, '10.1 Tạo và đổi tên gia đình', 2)
add_body(doc, 'Tài khoản mới chưa thuộc gia đình sẽ được chuyển đến màn hình Tạo gia đình. Nhập tên gia đình để bắt đầu; người tạo trở thành Chủ gia đình.')
add_bullet(doc, 'Chủ gia đình dùng biểu tượng bút chì ở màn hình Thành viên để đổi tên gia đình.')
add_bullet(doc, 'Chỉ xóa được gia đình không còn giao dịch đang hoạt động. Các giao dịch đã xóa mềm sẽ được dọn khi xóa gia đình và không thể khôi phục.')
add_callout(doc, 'Trước khi xóa gia đình', 'Hãy kiểm tra đúng gia đình và xuất dữ liệu nếu cần lưu trữ. Nếu còn giao dịch đang hoạt động, hệ thống sẽ chặn xóa.', 'warning')
add_heading(doc, '10.2 Thêm và quản lý thành viên', 2)
add_step(doc, 1, 'Yêu cầu người mới đăng ký', 'email phải có tài khoản Family Expense trước.')
add_step(doc, 2, 'Mở “Thành viên”', 'chủ gia đình nhập email và tên hiển thị, sau đó thêm vào gia đình.')
add_step(doc, 3, 'Đăng nhập lại', 'thành viên sẽ thấy chung Dashboard, danh mục và giao dịch của gia đình.')
add_bullet(doc, 'Chủ gia đình có thể đổi tên hiển thị của mọi thành viên và xóa thành viên khỏi gia đình.')
add_bullet(doc, 'Thành viên có thể đổi tên của chính mình, nhưng không thể thêm người khác hoặc quản lý danh mục.')
add_bullet(doc, 'Xóa thành viên chỉ thu hồi quyền truy cập; giao dịch người đó đã tạo vẫn được giữ trong lịch sử.')
add_callout(doc, 'Phạm vi hiện tại', 'Mỗi tài khoản chỉ thuộc một gia đình đang hoạt động tại một thời điểm.')

add_heading(doc, '11. Cài ứng dụng trên điện thoại', 1)
add_heading(doc, '11.1 iPhone/iPad (Safari)', 2)
add_step(doc, 1, 'Mở bằng Safari', 'truy cập domain Family Expense và đăng nhập.')
add_step(doc, 2, 'Mở menu Chia sẻ', 'chọn biểu tượng hình vuông có mũi tên hướng lên.')
add_step(doc, 3, 'Chọn “Thêm vào Màn hình chính”', 'xác nhận tên rồi chọn “Thêm”.')
add_body(doc, 'Biểu tượng Family Expense sẽ xuất hiện như một ứng dụng độc lập. Khi có phiên bản mới, đóng mở lại app hoặc tải lại trang để cập nhật cache.')
add_heading(doc, '11.2 Android (Chrome)', 2)
add_step(doc, 1, 'Mở bằng Chrome', 'truy cập Family Expense.')
add_step(doc, 2, 'Mở menu trình duyệt', 'chọn “Cài đặt ứng dụng” hoặc “Thêm vào màn hình chính”.')
add_step(doc, 3, 'Xác nhận cài đặt', 'mở app từ biểu tượng mới trên màn hình chính.')
add_callout(doc, 'Khi ngoại tuyến', 'Ứng dụng có thể mở màn hình ngoại tuyến, nhưng giao dịch chỉ được xem là đã lưu khi có mạng và database xác nhận thành công.', 'warning')

add_heading(doc, '12. Quy tắc dữ liệu quan trọng', 1)
rules = [
    ('Số tiền', 'Luôn là số dương. Loại giao dịch quyết định khoản đó là thu hay chi.'),
    ('Ngày', 'Quyết định kỳ tháng/năm và trạng thái mặc định.'),
    ('Trạng thái', 'Chỉ giao dịch Thực tế được dùng cho KPI và phần tổng hợp thực tế.'),
    ('Dự kiến', 'Khi đến hạn cần người dùng xác nhận; hệ thống không tự chuyển im lặng.'),
    ('Phương thức', 'Bắt buộc; mặc định Chuyển khoản khi tạo mới.'),
    ('Chi ròng', 'Chi tiêu mang dấu cộng; Thu nhập mang dấu trừ trên tổng danh sách.'),
    ('Xóa', 'Giao dịch đã xóa không xuất hiện trong danh sách/file xuất; member chỉ xóa được dòng do mình tạo.'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
set_table_widths(table, [1.6, 4.9])
for i, text in enumerate(('Nội dung', 'Quy tắc')):
    shade(table.rows[0].cells[i], NAVY)
    set_font(table.rows[0].cells[i].paragraphs[0].add_run(text), bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
for label, detail in rules:
    cells = table.add_row().cells
    cells[0].text = label
    cells[1].text = detail
    set_font(cells[0].paragraphs[0].runs[0], bold=True, color=NAVY)
    set_font(cells[1].paragraphs[0].runs[0])
for row in table.rows:
    prevent_row_split(row)
    for cell in row.cells:
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

add_heading(doc, '13. Mẹo sử dụng và xử lý sự cố', 1)
add_heading(doc, '13.1 Thói quen nên áp dụng', 2)
add_bullet(doc, 'Ghi giao dịch ngay sau khi phát sinh để tránh quên.')
add_bullet(doc, 'Mỗi tuần kiểm tra danh sách Dự kiến đến hạn.')
add_bullet(doc, 'Mỗi tháng đối chiếu KPI với sao kê ngân hàng và xuất một bản Excel lưu trữ.')
add_bullet(doc, 'Dùng nội dung dễ tìm, ví dụ “Tiền điện tháng 8” thay vì “Thanh toán”.')
add_heading(doc, '13.2 Một số vấn đề thường gặp', 2)
faq = [
    ('Không thấy thay đổi mới', 'Tải lại trang. Nếu đã cài dạng PWA, đóng ứng dụng và mở lại để nhận phiên bản mới.'),
    ('Không lưu được giao dịch', 'Kiểm tra các trường có dấu sao, số tiền phải lớn hơn 0 và thiết bị đang có mạng.'),
    ('AI phân tích chậm', 'Giữ câu nhập ngắn, mỗi câu một giao dịch; thử lại khi mạng ổn định. Bạn vẫn có thể tiếp tục điền form thủ công mà không cần đổi tab.'),
    ('AI chọn sai danh mục', 'Chọn lại danh mục trước khi lưu và cân nhắc đổi tên danh mục cho rõ nghĩa hơn.'),
    ('Không xóa được danh mục', 'Danh mục có thể đã được dùng. Hãy giữ lại để bảo toàn lịch sử giao dịch.'),
    ('File Excel bị báo sai template', 'Tải lại template từ màn hình Dữ liệu; không đổi tên sheet “Giao dịch”, tiêu đề cột hoặc định dạng .xlsx.'),
    ('Không thấy dữ liệu gia đình', 'Kiểm tra tài khoản đã được chủ gia đình thêm đúng email; sau đó đăng xuất và đăng nhập lại.'),
    ('Không xóa được giao dịch của người khác', 'Member chỉ xóa giao dịch do chính mình tạo; hãy liên hệ chủ gia đình nếu cần xử lý.'),
    ('Không thấy nút sửa danh mục', 'Tài khoản hiện tại có thể là thành viên; chỉ chủ gia đình được quản lý danh mục.'),
    ('Số liệu dashboard không đúng kỳ', 'Kiểm tra cả Tháng và Năm, đồng thời xác nhận giao dịch đã ở trạng thái Thực tế.'),
]
for question, answer in faq:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(question)
    set_font(r, bold=True, color=NAVY)
    p2 = doc.add_paragraph(answer)
    p2.paragraph_format.left_indent = Inches(0.18)
    p2.paragraph_format.keep_together = True
    p2.paragraph_format.space_after = Pt(3)
    p2.paragraph_format.line_spacing = 1.1
    set_font(p2.runs[0], size=10.5)

add_heading(doc, 'Checklist cuối tháng', 2)
for item in [
    'Chọn đúng tháng và năm trên Tổng quan.',
    'Xác nhận các giao dịch Dự kiến đã phát sinh.',
    'Đối chiếu Tổng thu, Tổng chi và Giá trị ròng.',
    'Tìm và xử lý giao dịch trùng hoặc sai thông tin.',
    'Xuất file Excel và lưu ở nơi an toàn.',
]:
    p = add_bullet(doc, item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1

doc.core_properties.title = 'Hướng dẫn sử dụng Family Expense'
doc.core_properties.subject = 'Tài liệu hướng dẫn dành cho người dùng cuối'
doc.core_properties.author = 'Family Expense'
doc.core_properties.keywords = 'Family Expense, hướng dẫn, giao dịch, quản lý chi tiêu'
doc.core_properties.comments = 'Cập nhật theo phiên bản ứng dụng ngày 27/08/2026.'
doc.save(OUT)
print(OUT)
