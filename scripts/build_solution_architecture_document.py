from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Solution_Architecture_Document_Family_Expense.docx"

NAVY = "163C35"
GREEN = "0F6B50"
MINT = "E8F4EF"
LIGHT = "F3F6F5"
GRAY = "5F6B7A"
RED = "9F1239"
AMBER = "8A4B08"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_widths(table, widths_inches):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        dxa = int(width * 1440)
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(dxa))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            dxa = int(widths_inches[idx] * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(9)
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            if ridx % 2 == 1:
                set_cell_shading(cells[i], LIGHT)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            r.font.size = Pt(8.5)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)


def add_numbered(doc, items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids or [0]) + 1
    num_id = max(num_ids or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720"); tabs.append(tab)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360")
    p_pr.append(tabs); p_pr.append(ind)
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr]); abstract.append(lvl); numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    aid = OxmlElement("w:abstractNumId"); aid.set(qn("w:val"), str(abstract_id)); num.append(aid); numbering.append(num)
    for item in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
        nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, nid]); p_pr.append(num_pr)
        p.add_run(item)


def add_callout(doc, label, text, fill=MINT, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(text)
    set_table_widths(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7F6")
    set_cell_margins(cell, 120, 140, 120, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(code.strip())
    r.font.name = "Courier New"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    r.font.size = Pt(7.3)
    set_table_widths(table, [6.5])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 11.5, 8, 4)):
        s = styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(GREEN if name != "Heading 3" else NAVY)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.text = "FAMILY EXPENSE  |  SOLUTION ARCHITECTURE DOCUMENT"
    header.style = styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Nội bộ - Phiên bản 1.0 | 28/08/2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(GRAY)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    r = p.add_run("SOLUTION ARCHITECTURE\nDOCUMENT")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p = doc.add_paragraph()
    r = p.add_run("Family Expense")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph("Nền tảng quản lý giao dịch gia đình tiếng Việt, mobile-first và PWA")
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    p.paragraph_format.space_after = Pt(28)
    add_table(doc, ["Thuộc tính", "Giá trị"], [
        ("Phiên bản", "1.0"),
        ("Ngày ban hành", "28/08/2026 - Asia/Ho_Chi_Minh"),
        ("Trạng thái", "Baseline kiến trúc hiện tại và kiến trúc mục tiêu"),
        ("Đối tượng", "Ban lãnh đạo, Product Owner, đội kỹ thuật, vận hành và bảo mật"),
        ("Phạm vi", "Frontend, Supabase backend, Gemini AI, dữ liệu Excel và Cloudflare Pages"),
    ], [1.45, 5.05])
    add_callout(doc, "Khuyến nghị kiến trúc", "Giữ mô hình serverless hiện tại cho MVP; ưu tiên thiết lập Git, CI/CD, staging và giám sát trước khi mở rộng thêm tính năng.")
    doc.add_page_break()

    heading(doc, "Mục lục", 1)
    for item in [
        "1. Executive Summary", "2. Business Context & Goals", "3. Scope & Assumptions",
        "4. Current State Architecture", "5. Proposed Solution Architecture",
        "6. Technology Stack & Justification", "7. Non-Functional Requirements",
        "8. Deployment Architecture", "9. Security Architecture",
        "10. Risks, Assumptions, Dependencies, Constraints (RAID)",
        "11. Migration Strategy", "12. Roadmap / Phasing", "13. Appendix",
    ]:
        doc.add_paragraph(item)
    doc.add_page_break()

    heading(doc, "1. Executive Summary", 1)
    doc.add_paragraph("Family Expense là ứng dụng web/PWA quản lý giao dịch gia đình bằng tiếng Việt, thay thế quy trình Excel thủ công và tạo một nguồn dữ liệu thống nhất cho nhiều thành viên. Kiến trúc hiện tại sử dụng React/TypeScript trên Cloudflare Pages và Supabase Cloud cho xác thực, PostgreSQL, Row Level Security (RLS), RPC và Edge Functions. Gemini được gọi duy nhất từ Edge Function để phân tích câu tiếng Việt thành đề xuất có cấu trúc; người dùng luôn kiểm tra trước khi lưu.")
    doc.add_paragraph("Giải pháp phù hợp với giai đoạn MVP và ngân sách thấp nhờ dịch vụ managed/serverless có Free Tier. Điểm mạnh là ranh giới bảo mật ở database, triển khai frontend nhẹ và luồng AI có human-in-the-loop. Rủi ro vận hành lớn nhất hiện nay là thiếu Git repository, CI/CD và môi trường staging chuẩn hóa; deploy thủ công làm tăng nguy cơ sai khác phiên bản và khó rollback.")
    add_callout(doc, "Quyết định đề xuất", "Không tái kiến trúc nền tảng. Chuẩn hóa delivery pipeline, phân tách staging/production, bổ sung observability và kiểm thử RLS trước khi mở rộng Storage, giao dịch định kỳ hoặc tự động hóa AI.")

    heading(doc, "2. Business Context & Goals", 1)
    heading(doc, "2.1 Bối cảnh kinh doanh", 2)
    doc.add_paragraph("Dữ liệu chi tiêu gia đình trước đây được nhập, phân loại và tổng hợp bằng Excel. Cách làm này khó cộng tác, dễ trùng dữ liệu, thiếu kiểm soát truy cập theo gia đình và không thuận tiện trên điện thoại. Family Expense chuyển quy trình sang một ứng dụng dùng chung, giữ khả năng import/export để giảm gián đoạn thay đổi thói quen.")
    heading(doc, "2.2 Mục tiêu", 2)
    add_bullets(doc, [
        "Quản lý tập trung giao dịch tiền vào và tiền ra theo gia đình, thành viên, mục đích, danh mục và phương thức.",
        "Cung cấp dashboard, tìm kiếm, lọc, sửa hàng loạt, thùng rác và export phục vụ đối soát.",
        "Hỗ trợ nhập nhanh bằng tiếng Việt qua Gemini nhưng không tự động ghi dữ liệu.",
        "Cài đặt như PWA, dùng tốt bằng một tay trên điện thoại và responsive trên desktop.",
        "Giữ chi phí vận hành thấp, tránh quản lý máy chủ và bảo vệ dữ liệu bằng RLS.",
    ])
    heading(doc, "2.3 Stakeholders", 2)
    add_table(doc, ["Nhóm", "Mối quan tâm", "Trách nhiệm"], [
        ("Chủ gia đình (Owner)", "Độ chính xác, quyền quản trị, dữ liệu riêng tư", "Quản lý thành viên, danh mục và dữ liệu"),
        ("Thành viên", "Nhập nhanh, xem đúng dữ liệu gia đình", "Tạo và cập nhật giao dịch được phép"),
        ("Product Owner", "Giá trị sử dụng, phạm vi MVP, chi phí", "Ưu tiên backlog và nghiệm thu"),
        ("Kỹ thuật", "Khả năng bảo trì, test, deploy, bảo mật", "Phát triển, migration, vận hành"),
        ("Nhà cung cấp", "Quota, SLA, thay đổi API", "Cloudflare, Supabase, Google AI"),
    ], [1.45, 2.35, 2.70])

    heading(doc, "3. Scope & Assumptions", 1)
    add_table(doc, ["In-scope", "Out-of-scope hiện tại"], [
        ("Auth email/password, reset/magic link; gia đình và thành viên", "OCR hóa đơn, hình ảnh và giọng nói phía server"),
        ("CRUD, soft delete, khôi phục, xóa vĩnh viễn, bulk edit", "Open banking, đồng bộ ngân hàng tự động"),
        ("Dashboard, tìm kiếm không dấu, lọc, phân trang/tải thêm", "Kế toán doanh nghiệp, thuế, đa tiền tệ nâng cao"),
        ("Danh mục, mục đích, phương thức, tài khoản, đối tượng", "Workflow phê duyệt nhiều cấp"),
        ("Import nguồn Excel, template, export đầy đủ", "Data warehouse/BI độc lập"),
        ("AI text-to-structured suggestion và PWA", "AI tự lưu hoặc tự quyết định giao dịch"),
    ], [3.25, 3.25])
    heading(doc, "3.1 Giả định chính", 2)
    add_bullets(doc, [
        "Mỗi request dữ liệu được gắn family_id và người dùng đã xác thực.",
        "VND là tiền tệ mặc định; amount lưu số nguyên dương, không có phần thập phân.",
        "Múi giờ nghiệp vụ là Asia/Ho_Chi_Minh; ngày hiển thị dd/MM/yyyy.",
        "Nhãn UI Tiền ra/Tiền vào ánh xạ với enum database Chi tiêu/Thu nhập; dữ liệu legacy đã được chuẩn hóa trước khi loại bỏ khỏi schema.",
        "Storage mới chỉ là điểm mở rộng kiến trúc; MVP không lưu chứng từ.",
        "Free Tier phù hợp quy mô gia đình/MVP nhưng không được coi là cam kết SLA production.",
    ])

    heading(doc, "4. Current State Architecture", 1)
    doc.add_paragraph("Hiện trạng là kiến trúc SPA/PWA serverless. Trình duyệt giao tiếp trực tiếp với Supabase bằng anon key; JWT và RLS quyết định quyền dữ liệu. Các thao tác phức tạp/atomic dùng PostgreSQL RPC. AI đi qua Edge Function để bảo vệ Gemini API key.")
    add_code(doc, """flowchart LR
  U[Người dùng Web/PWA] --> CF[Cloudflare Pages CDN]
  CF --> SPA[React 19 SPA]
  SPA --> AUTH[Supabase Auth]
  SPA --> API[PostgREST / RPC]
  API --> DB[(PostgreSQL + RLS)]
  SPA --> EF[Edge Function parse-expense]
  EF --> DB
  EF --> GEM[Gemini API]
  SPA -. import/export .-> XLSX[Excel/CSV cục bộ]""")
    heading(doc, "4.1 Đánh giá hiện trạng", 2)
    add_table(doc, ["Khía cạnh", "Hiện trạng", "Đánh giá"], [
        ("Ứng dụng", "React/Vite PWA trên Cloudflare Pages", "Phù hợp MVP, tải nhanh, vận hành đơn giản"),
        ("Dữ liệu", "Supabase PostgreSQL, migrations, RPC, RLS", "Tốt; cần test RLS tự động và backup drill"),
        ("AI", "Edge Function + Gemini structured output", "Đúng ranh giới secret; phụ thuộc quota/model"),
        ("Delivery", "Deploy thủ công; workspace chưa có Git", "Rủi ro cao về truy vết, rollback và tái lập"),
        ("Môi trường", "Local và production; staging chưa chuẩn", "Cần tách project/dataset staging"),
        ("Observability", "Log nền tảng và ai_usage_logs", "Cần alert, dashboard lỗi và retention rõ ràng"),
    ], [1.1, 2.75, 2.65])

    heading(doc, "5. Proposed Solution Architecture", 1)
    heading(doc, "5.1 Nguyên tắc kiến trúc", 2)
    add_bullets(doc, [
        "Managed-first/serverless-first; tránh thêm hạ tầng khi chưa có nhu cầu đo được.",
        "Database là security boundary; frontend validation không thay thế RLS/constraint/RPC.",
        "Human-in-the-loop cho AI; structured output và allow-list danh mục.",
        "Idempotency và audit cho import, bulk mutation và thao tác xóa.",
        "Tách cấu hình/môi trường, triển khai bất biến và rollback được.",
        "Không cache quá mức dữ liệu tài chính trên service worker.",
    ])
    heading(doc, "5.2 Kiến trúc tổng quan mục tiêu", 2)
    add_code(doc, """flowchart TB
  subgraph Client[Client Layer]
    PWA[React PWA] --> UI[Router + Forms + Query Cache]
  end
  subgraph Delivery[Delivery & Edge]
    GH[GitHub] --> CI[CI: lint/typecheck/test/build]
    CI --> STG[Cloudflare Pages Staging]
    CI --> PRD[Cloudflare Pages Production]
  end
  subgraph Supabase[Supabase per Environment]
    AUTH[Auth] --> JWT[JWT]
    REST[PostgREST] --> RLS[RLS Policies]
    RPC[Transactional RPC] --> PG[(PostgreSQL)]
    EF[parse-expense Edge Function] --> PG
  end
  PWA --> AUTH
  UI --> REST
  UI --> RPC
  UI --> EF
  EF --> GEM[Google Gemini Flash]
  OBS[Logs / Alerts / Audit] --- PRD
  OBS --- Supabase""")
    heading(doc, "5.3 Component breakdown", 2)
    add_table(doc, ["Thành phần", "Trách nhiệm", "Ranh giới/ghi chú"], [
        ("React PWA", "UI, routing, form, accessibility, offline shell", "Không chứa service-role/Gemini key; không xác nhận lưu khi request chưa thành công"),
        ("TanStack Query/App Context", "Server state, invalidation, session refresh", "Cache ngắn hạn; refresh session khi app foreground"),
        ("Supabase Auth", "Đăng ký, đăng nhập, reset, JWT", "Redirect URL theo môi trường; xử lý clock skew và refresh token"),
        ("PostgREST + RLS", "CRUD dữ liệu gia đình", "Mọi bảng family-scoped bật RLS"),
        ("PostgreSQL RPC", "Onboarding, import atomic, bulk update/delete, tổng hợp", "SECURITY DEFINER tối thiểu; validate auth.uid/family"),
        ("Edge Function", "Xác thực, rate limit, prompt, validate AI", "Không log nội dung nhạy cảm; lỗi 429 không retry vô hạn"),
        ("Gemini", "Phân tích một câu thành JSON", "Model qua env; chỉ chọn danh mục hiện có"),
        ("Cloudflare Pages", "Static hosting, CDN, SPA fallback", "Preview/staging và production tách biệt"),
    ], [1.3, 2.55, 2.65])
    heading(doc, "5.4 Data architecture", 2)
    doc.add_paragraph("Mô hình dữ liệu lấy family_id làm tenant key. families/family_members xác định tenant và role; transactions là fact chính; purposes, expense_types (hiển thị là Danh mục), payment_methods, accounts, events và beneficiaries là dimensions. budgets và recurring_transactions tồn tại ở schema nhưng chức năng quản lý nâng cao không nằm trong MVP hiện tại.")
    add_code(doc, """erDiagram
  FAMILIES ||--o{ FAMILY_MEMBERS : contains
  FAMILIES ||--o{ TRANSACTIONS : owns
  FAMILIES ||--o{ PURPOSES : configures
  FAMILIES ||--o{ EXPENSE_TYPES : configures
  FAMILIES ||--o{ PAYMENT_METHODS : configures
  PURPOSES ||--o{ TRANSACTIONS : classifies
  EXPENSE_TYPES ||--o{ TRANSACTIONS : categorizes
  PAYMENT_METHODS ||--o{ TRANSACTIONS : pays
  EVENTS ||--o{ TRANSACTIONS : groups
  BENEFICIARIES ||--o{ TRANSACTIONS : benefits
  ACCOUNTS ||--o{ TRANSACTIONS : uses""")
    heading(doc, "5.5 Luồng ghi giao dịch", 2)
    add_code(doc, """sequenceDiagram
  actor U as Người dùng
  participant P as React PWA
  participant S as Supabase API
  participant D as PostgreSQL/RLS
  U->>P: Nhập/chỉnh sửa và xác nhận
  P->>P: Validate Zod + kiểm tra số tiền
  P->>S: Mutation kèm JWT + family_id
  S->>D: RLS/constraint/RPC
  alt Thành công
    D-->>P: Bản ghi đã commit
    P->>P: Invalidate query và hiển thị thành công
  else Thất bại
    D-->>P: Lỗi có mã
    P-->>U: Giữ form, hiển thị lỗi; không giả định đã lưu
  end""")
    heading(doc, "5.6 Luồng AI", 2)
    add_code(doc, """sequenceDiagram
  actor U as Người dùng
  participant P as PWA
  participant E as parse-expense
  participant D as Supabase DB
  participant G as Gemini
  U->>P: Nhập câu tiếng Việt, bấm AI
  P->>E: text + familyId + timezone + JWT
  E->>D: Kiểm tra membership, quota, lấy danh mục
  E->>G: Prompt + JSON Schema + allow-list
  G-->>E: Structured JSON
  E->>E: Zod + ID allow-list validation
  E-->>P: Suggestion + confidence + warnings
  P-->>U: Tô sáng trường AI đã điền để kiểm tra
  U->>P: Chỉnh sửa và xác nhận
  P->>D: Lưu transaction source=ai
  Note over P,D: AI không bao giờ tự lưu""")
    heading(doc, "5.7 Integration architecture", 2)
    add_table(doc, ["Tích hợp", "Giao thức", "Kiểm soát"], [
        ("PWA ↔ Supabase", "HTTPS/JSON, PostgREST, RPC", "JWT, RLS, Zod, family_id, timeout/error state"),
        ("Edge Function ↔ Gemini", "HTTPS GenerateContent", "Secret server-side, model env, JSON Schema, rate limit"),
        ("PWA ↔ Excel", "XLSX xử lý phía client/CLI", "Schema/mapping, dry-run, duplicate check, batch audit"),
        ("CI ↔ Cloudflare", "API token / official action", "Secret scope nhỏ, preview trước production, approval"),
        ("CI ↔ Supabase", "Supabase CLI", "Migration additive, linked project theo environment, approval production"),
    ], [1.55, 1.75, 3.20])
    doc.add_paragraph("Message queue chưa cần thiết ở quy mô hiện tại. Khi import hoặc AI vượt ngưỡng đồng bộ (ví dụ >30 giây, hàng chục nghìn dòng), cân nhắc Supabase Queues/Cloudflare Queues và job status; không đưa queue vào MVP chỉ để dự phòng.")

    heading(doc, "6. Technology Stack & Justification", 1)
    add_table(doc, ["Lớp", "Lựa chọn", "Lý do", "Phương án khác"], [
        ("Frontend", "React 19 + TypeScript + Vite", "Phổ biến, strict typing, build nhanh, phù hợp SPA/PWA", "Next.js: thừa SSR cho app đăng nhập; tăng vận hành"),
        ("UI", "Tailwind + RHF + Zod", "Mobile-first, form hiệu quả, schema dùng lại", "Component suite lớn: nhanh ban đầu nhưng tăng bundle/style lock-in"),
        ("Server state", "TanStack Query", "Cache/invalidation/mutation state chuẩn", "Context thuần: khó scale và dễ stale"),
        ("Database/BaaS", "Supabase PostgreSQL", "Auth, RLS, SQL/RPC, migration trong một nền tảng", "Firebase: NoSQL và authorization kém phù hợp báo cáo quan hệ"),
        ("AI", "Gemini Flash qua Edge Function", "Structured output, chi phí thấp/Free Tier", "Model khác: cần đánh giá quota, latency, privacy"),
        ("Hosting", "Cloudflare Pages", "CDN toàn cầu, preview, SPA static phù hợp", "Vercel/Netlify tương đương; không tạo giá trị đổi nền tảng"),
        ("Testing", "Vitest/RTL/Playwright", "Bao phủ logic, component và luồng chính", "Chỉ E2E: chậm và khó chẩn đoán"),
    ], [1.0, 1.55, 2.55, 1.40])

    heading(doc, "7. Non-Functional Requirements", 1)
    add_table(doc, ["Thuộc tính", "Mục tiêu đề xuất", "Cách đo/kiểm soát"], [
        ("Hiệu năng", "P95 tải route chính <2,5s trên 4G; query giao dịch P95 <1s ở 10.000 bản ghi/family", "Web Vitals, Supabase query stats, index family/date và search"),
        ("Khả năng mở rộng", "Tăng theo số family mà không đổi kiến trúc; tải thêm theo trang", "Keyset/range pagination, RPC aggregate, giới hạn bulk 100/lần"),
        ("Sẵn sàng", "Mục tiêu nội bộ 99,5%/tháng, phụ thuộc SLA nhà cung cấp", "Status page, synthetic health check, runbook"),
        ("Bảo mật", "Không truy cập chéo family; không lộ secret", "RLS tests, bundle scan, secret manager, least privilege"),
        ("Khả năng phục hồi", "RPO 24h, RTO 4h cho MVP", "PITR/backup theo gói, export định kỳ, restore drill hàng quý"),
        ("Usability", "Mobile-first, keyboard/label/contrast/loading/empty nhất quán", "RTL, Playwright mobile viewport, accessibility checks"),
        ("Maintainability", "Build tái lập, migration bất biến, changelog", "CI gates, review, ADR, dependency audit"),
        ("Privacy", "Tối thiểu hóa log và dữ liệu gửi AI", "Không log token/input tài chính; retention ai_usage_logs"),
    ], [1.2, 2.65, 2.65])

    heading(doc, "8. Deployment Architecture", 1)
    heading(doc, "8.1 Môi trường mục tiêu", 2)
    add_table(doc, ["Môi trường", "Frontend", "Backend/Data", "Mục đích"], [
        ("Local", "Vite localhost", "Supabase local hoặc test project", "Phát triển, unit/component test"),
        ("Preview", "Cloudflare Pages PR preview", "Không mutation production", "Review UI/build theo commit"),
        ("Staging", "Domain staging cố định", "Supabase staging riêng, dữ liệu giả lập", "E2E, RLS, migration rehearsal"),
        ("Production", "family-expense-8fo.pages.dev/custom domain", "Supabase production", "Người dùng thật"),
    ], [1.0, 1.75, 2.15, 1.60])
    heading(doc, "8.2 CI/CD đề xuất", 2)
    add_code(doc, """flowchart LR
  DEV[Feature branch] --> PR[Pull Request]
  PR --> Q[lint + typecheck + unit + build]
  Q --> PRE[Cloudflare Preview]
  PRE --> REV[Review/Approve]
  REV --> MAIN[Merge main]
  MAIN --> STG[Deploy Staging]
  STG --> E2E[E2E + RLS + migration smoke]
  E2E --> GATE[Manual production approval]
  GATE --> MIG[Supabase db push]
  MIG --> EDGE[Deploy Edge Function]
  EDGE --> PROD[Cloudflare Production]
  PROD --> SMOKE[Smoke test + monitor + rollback if needed]""")
    add_bullets(doc, [
        "Khởi tạo private GitHub repository; bảo vệ main, bắt buộc review và CI xanh.",
        "Dùng environment secrets của GitHub/Cloudflare/Supabase; không lưu secret trong repo.",
        "Migration production chạy trước code chỉ khi backward-compatible; thay đổi phá vỡ dùng expand-migrate-contract.",
        "Cloudflare rollback bằng deployment trước; database rollback ưu tiên forward-fix migration.",
        "Gắn version/commit SHA vào build metadata và release note.",
    ])

    heading(doc, "9. Security Architecture", 1)
    heading(doc, "9.1 AuthN/AuthZ", 2)
    add_bullets(doc, [
        "Supabase Auth phát JWT; frontend duy trì session và chủ động refresh khi PWA quay lại foreground.",
        "RLS kiểm tra membership active theo family_id; owner quản trị thành viên/danh mục, member thao tác giao dịch theo policy/RPC.",
        "Edge Function yêu cầu Bearer JWT, gọi RPC get_ai_request_context để kiểm tra membership và quota.",
        "Không dùng service-role key ở frontend; anon key chỉ an toàn khi RLS/privilege được cấu hình đúng.",
    ])
    heading(doc, "9.2 Bảo vệ dữ liệu và secret", 2)
    add_table(doc, ["Vùng", "Kiểm soát bắt buộc"], [
        ("Truyền dữ liệu", "TLS/HTTPS; HSTS/custom domain khi cấu hình"),
        ("Dữ liệu lưu", "Mã hóa managed của Supabase; không lưu đầy đủ số thẻ/tài khoản"),
        ("Secret", "Supabase Secrets/GitHub Environment Secrets; rotation và scope tối thiểu"),
        ("Logs", "Không token, API key, nội dung AI nhạy cảm; chỉ metadata vận hành"),
        ("PWA cache", "Cache application shell; không cache lâu response tài chính hoặc mutation"),
        ("Input", "Zod ở biên, DB constraints, allow-list IDs, giới hạn kích thước"),
        ("Supply chain", "Lockfile, dependency review, SCA/secret scan trong CI"),
    ], [1.45, 5.05])
    heading(doc, "9.3 Threat model rút gọn", 2)
    add_table(doc, ["Mối đe dọa", "Giảm thiểu"], [
        ("IDOR/truy cập chéo gia đình", "RLS + family_id scope + negative security tests"),
        ("JWT cũ/clock skew trên PWA", "Auto refresh foreground, retry có giới hạn sau refresh, thông báo đăng nhập lại"),
        ("Prompt injection/danh mục bịa", "Structured output, prompt giới hạn, server-side allow-list ID"),
        ("Import độc hại/trùng", "Schema validation, dry-run, giới hạn file/dòng, unique source reference, audit batch"),
        ("Xóa ngoài ý muốn", "Soft delete, xác nhận, quyền owner cho xóa vĩnh viễn, audit"),
        ("Lộ secret qua bundle/log", "Bundle scan, secret server-side, log redaction"),
    ], [2.05, 4.45])
    doc.add_paragraph("Compliance: ứng dụng chưa tuyên bố tuân thủ PCI DSS hoặc chuẩn kế toán. Thiết kế giảm phạm vi PCI bằng cách chỉ lưu tên tổ chức và bốn số cuối, không lưu PAN/CVV. Nếu phục vụ khách hàng bên ngoài, cần đánh giá chính sách riêng tư, quyền xóa/xuất dữ liệu và quy định địa phương.")

    heading(doc, "10. Risks, Assumptions, Dependencies, Constraints (RAID)", 1)
    add_table(doc, ["Loại", "Mục", "Mức", "Ứng phó/Chủ trì"], [
        ("Risk", "Chưa có Git/CI/CD/staging", "Cao", "Thiết lập repository, branch protection, pipeline và staging trước feature mới - Tech Lead"),
        ("Risk", "Free Tier quota/rate limit", "Cao", "Theo dõi quota, graceful error, ngân sách nâng cấp - Product/Tech"),
        ("Risk", "RLS sai gây lộ chéo family", "Rất cao", "Automated negative tests và review migration - Backend/Security"),
        ("Risk", "Model Gemini đổi/deprecate", "Trung bình", "Model qua env, contract test, fallback runbook - Tech Lead"),
        ("Risk", "JWT issued-at-future trên mobile", "Trung bình", "Refresh foreground, clock guidance, telemetry không nhạy cảm - Frontend"),
        ("Assumption", "Quy mô MVP nhỏ, tải đồng bộ đủ dùng", "-", "Đo P95 và số dòng trước khi thêm queue"),
        ("Dependency", "Supabase/Cloudflare/Google AI", "Cao", "Theo dõi status/SLA, export và runbook sự cố"),
        ("Constraint", "VND, tiếng Việt, Asia/Ho_Chi_Minh", "Cố định", "Centralize formatter/domain constants"),
        ("Constraint", "AI không tự lưu", "Cố định", "UI confirmation + backend không có auto-insert path"),
        ("Constraint", "Không phá dữ liệu enum cũ", "Cố định", "Mapping label mới, migration tương thích"),
    ], [0.75, 2.45, 0.72, 2.58])

    heading(doc, "11. Migration Strategy", 1)
    heading(doc, "11.1 Migration dữ liệu Excel", 2)
    add_numbered(doc, [
        "Đóng băng workbook nguồn; lưu checksum và không sửa file gốc.",
        "Chạy dry-run, kiểm tra sheet Giao dịch chuẩn hóa và mapping danh mục.",
        "Tạo/seed family, mục đích, danh mục và phương thức trước giao dịch.",
        "Import theo batch bằng RPC transaction; dòng lỗi được cô lập, không làm hỏng toàn batch.",
        "Chống trùng bằng source_reference hoặc fingerprint ngày + nội dung chuẩn hóa + số tiền.",
        "Đối chiếu 2.090 giao dịch và tổng ròng 1.696.313.649 VND; ghi import_batches/audit.",
        "Business owner duyệt danh sách ngoại lệ rồi mới đóng migration.",
    ])
    heading(doc, "11.2 Migration kiến trúc vận hành", 2)
    add_numbered(doc, [
        "Khởi tạo Git từ workspace hiện tại, loại trừ .env/node_modules/.supabase temp; tạo baseline tag.",
        "Tạo Supabase staging và Cloudflare staging; sao chép schema bằng migrations, không dùng dữ liệu tài chính thật.",
        "Đưa quality gates vào CI; chạy build và test trên pull request.",
        "Chuyển deploy production sang pipeline có phê duyệt; xác minh URL, auth redirect và smoke test.",
        "Thực hiện rollback rehearsal và backup restore drill trước khi coi pipeline là hoàn tất.",
    ])

    heading(doc, "12. Roadmap / Phasing", 1)
    add_table(doc, ["Giai đoạn", "Thời lượng gợi ý", "Kết quả", "Exit criteria"], [
        ("0. Baseline", "1-2 ngày", "Tag phiên bản, tài liệu SAD/ADR, inventory secret", "Build tái lập; owner xác nhận baseline"),
        ("1. Delivery foundation", "2-4 ngày", "GitHub, CI, preview, branch protection", "PR bắt buộc CI xanh; rollback frontend được"),
        ("2. Staging & DB safety", "3-5 ngày", "Supabase staging, migration rehearsal, RLS tests", "Không truy cập chéo family; db push staging thành công"),
        ("3. Observability", "2-3 ngày", "Error tracking, uptime, alert/quota dashboard, runbook", "Cảnh báo thử nghiệm đến đúng người"),
        ("4. Hardening", "3-5 ngày", "E2E auth/import/AI, backup drill, security review", "RPO/RTO được kiểm chứng"),
        ("5. Feature evolution", "Theo backlog", "Storage chứng từ, recurring, queue khi có số liệu", "Có business case và NFR đo được"),
    ], [1.15, 1.05, 2.25, 2.05])
    add_callout(doc, "Ưu tiên 30 ngày", "Hoàn tất giai đoạn 0-3 trước khi thêm tích hợp mới. Đây là thay đổi nâng độ tin cậy lớn nhất mà không làm tăng đáng kể độ phức tạp sản phẩm.")

    heading(doc, "13. Appendix", 1)
    heading(doc, "13.1 Glossary", 2)
    add_table(doc, ["Thuật ngữ", "Định nghĩa"], [
        ("RLS", "Row Level Security - policy PostgreSQL giới hạn từng hàng theo user/family."),
        ("RPC", "Hàm PostgreSQL gọi từ API để xử lý nghiệp vụ atomic."),
        ("PWA", "Progressive Web App - web có manifest, service worker và khả năng cài đặt."),
        ("RPO/RTO", "Mức mất dữ liệu tối đa / thời gian phục hồi mục tiêu."),
        ("Human-in-the-loop", "AI đề xuất; con người kiểm tra và xác nhận quyết định cuối."),
        ("Expand-migrate-contract", "Thay đổi schema nhiều bước để tương thích ngược khi deploy."),
    ], [1.45, 5.05])
    heading(doc, "13.2 Decision records tóm tắt", 2)
    add_table(doc, ["ADR", "Quyết định", "Trạng thái"], [
        ("ADR-001", "Dùng Supabase làm backend managed và security boundary bằng RLS", "Accepted"),
        ("ADR-002", "Frontend SPA/PWA trên Cloudflare Pages", "Accepted"),
        ("ADR-003", "Gemini chỉ chạy qua Edge Function và không tự lưu", "Accepted"),
        ("ADR-004", "Giữ enum cũ, đổi nhãn UI Tiền vào/Tiền ra để tương thích", "Accepted - cần rà soát dài hạn"),
        ("ADR-005", "Thiết lập Git/CI/CD/staging trước feature lớn", "Proposed"),
    ], [0.9, 4.75, 0.85])
    heading(doc, "13.3 References", 2)
    add_bullets(doc, [
        "README.md và CHANGELOG.md của Family Expense (baseline 28/08/2026).",
        "Supabase migrations trong supabase/migrations và Edge Function parse-expense.",
        "Supabase Documentation: Auth, Row Level Security, Edge Functions và CLI.",
        "Cloudflare Pages Documentation: preview deployments, environment variables và rollback.",
        "Google Gemini API Documentation: structured output, model lifecycle và quota.",
        "OWASP ASVS và OWASP Top 10 cho kiểm soát ứng dụng web.",
    ])
    heading(doc, "13.4 Tiêu chí nghiệm thu kiến trúc", 2)
    add_bullets(doc, [
        "Không thể đọc/ghi dữ liệu của family khác ở cả PostgREST và RPC.",
        "Bundle frontend không chứa Gemini API key, service-role key hoặc secret deploy.",
        "AI trả structured suggestion, cảnh báo thiếu amount và không tạo bản ghi trước xác nhận.",
        "Import đối chiếu đúng 2.090 giao dịch và 1.696.313.649 VND, có danh sách ngoại lệ.",
        "CI chạy lint, typecheck, test và build; staging smoke test trước production.",
        "PWA không báo đã lưu khi mutation thất bại và refresh session sau foreground.",
        "Backup/restore và rollback deployment có runbook được thử nghiệm.",
    ])

    # keep headings with following content and avoid isolated rows where practical
    for table in doc.tables:
        for row in table.rows:
            prevent_row_split(row)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_together = True

    doc.core_properties.title = "Solution Architecture Document - Family Expense"
    doc.core_properties.subject = "Kiến trúc giải pháp Family Expense"
    doc.core_properties.author = "Family Expense Architecture Team"
    doc.core_properties.keywords = "Family Expense, SAD, React, Supabase, Cloudflare, Gemini, PWA"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
