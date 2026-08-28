from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from build_solution_architecture_document import (
    add_bullets, add_callout, add_code, add_numbered, add_table, heading,
    prevent_row_split, GREEN, NAVY, GRAY,
)

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Handoff_Document_Family_Expense.docx"


def checklist(doc, items):
    for status, text in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run("☒ " if status else "☐ ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(GREEN if status else "9F1239")
        p.add_run(text)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 11.5, 8, 4)):
        s = doc.styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(GREEN if name != "Heading 3" else NAVY)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    hp = sec.header.paragraphs[0]
    hp.text = "FAMILY EXPENSE  |  PROJECT HANDOFF DOCUMENT"
    hp.runs[0].font.size = Pt(8)
    hp.runs[0].font.bold = True
    hp.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Nội bộ - Phiên bản 1.0 | 28/08/2026")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run("PROJECT HANDOFF\nDOCUMENT")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p = doc.add_paragraph()
    r = p.add_run("Family Expense")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph("Tài liệu bàn giao sản phẩm, mã nguồn, dữ liệu và vận hành")
    add_table(doc, ["Thuộc tính", "Giá trị / trạng thái"], [
        ("Đội bàn giao", "TBD - cần điền tên đội/người phụ trách trước ký bàn giao"),
        ("Đội tiếp nhận", "TBD - cần điền tên đội/khách hàng trước ký bàn giao"),
        ("Ngày bàn giao", "28/08/2026 (dự thảo baseline)"),
        ("Trạng thái", "Ứng dụng đang chạy production; handoff vận hành chưa hoàn tất"),
        ("Production", "https://family-expense-8fo.pages.dev"),
        ("Phân loại", "Nội bộ - không chèn secret hoặc dữ liệu tài chính thật"),
    ], [1.55, 4.95])
    add_callout(doc, "Điều kiện ký nhận", "Hoàn thành các mục TBD về owner/contact, bàn giao quyền Cloudflare/Supabase/Google AI, thiết lập Git/CI/CD/staging và thực hiện restore/rollback drill.")
    doc.add_page_break()

    heading(doc, "Trạng thái bàn giao nhanh", 1)
    checklist(doc, [
        (True, "Ứng dụng production đang hoạt động trên Cloudflare Pages."),
        (True, "Schema, migrations, RLS/RPC và Edge Function nằm trong workspace."),
        (True, "README, CHANGELOG và SAD đã được cập nhật ở baseline hiện tại."),
        (False, "Chưa xác nhận Git repository/remote và chiến lược branch thực tế."),
        (False, "Chưa có staging và CI/CD production chuẩn hóa."),
        (False, "Chưa điền contact, owner alert, quyền truy cập và biên bản ký nhận."),
        (False, "Chưa kiểm chứng backup restore/rollback theo runbook."),
    ])

    heading(doc, "1. Project Overview", 1)
    add_table(doc, ["Nội dung", "Tóm tắt"], [
        ("Mục tiêu", "Thay Excel bằng ứng dụng quản lý giao dịch gia đình tiếng Việt, đa thành viên, PWA và nhập gợi ý bằng AI."),
        ("Phạm vi đã làm", "Auth, family/member, giao dịch, dashboard, danh mục, import/export, AI suggestion, bulk edit/delete/trash, PWA."),
        ("Ngoài phạm vi", "OCR, đồng bộ ngân hàng, kế toán doanh nghiệp, workflow phê duyệt, chứng từ Storage trong MVP."),
        ("Timeline hiện có", "Baseline sản phẩm và các thay đổi chính được ghi tại CHANGELOG đến 28/08/2026."),
        ("Trạng thái", "Go-live/đang chạy; tiếp tục hardening vận hành."),
    ], [1.45, 5.05])

    heading(doc, "2. Team & Contacts", 1)
    add_table(doc, ["Vai trò", "Tên", "Liên hệ", "Kênh", "Trạng thái"], [
        ("Product Owner", "TBD", "TBD", "TBD", "Cần điền"),
        ("Technical Owner", "TBD", "TBD", "TBD", "Cần điền"),
        ("Supabase Admin", "TBD", "TBD", "TBD", "Cần điền"),
        ("Cloudflare Admin", "TBD", "TBD", "TBD", "Cần điền"),
        ("Google AI Owner", "TBD", "TBD", "TBD", "Cần điền"),
        ("Incident Contact", "TBD", "TBD", "TBD", "Cần điền"),
    ], [1.25, 1.0, 1.55, 1.35, 1.35])
    add_callout(doc, "Quy tắc", "Không ghi mật khẩu, token hay API key trong tài liệu này. Chỉ ghi tên vault, nhóm quyền và quy trình cấp/thu hồi.", fill="FFF7E6", color="8A4B08")

    heading(doc, "3. Architecture Summary", 1)
    doc.add_paragraph("Frontend React/TypeScript là SPA/PWA trên Cloudflare Pages. Ứng dụng gọi Supabase Auth, PostgREST và RPC bằng anon key + JWT; PostgreSQL RLS là ranh giới phân quyền theo family. Edge Function parse-expense kiểm tra JWT/membership/rate limit rồi gọi Gemini bằng secret server-side và trả suggestion có cấu trúc để người dùng xác nhận.")
    add_code(doc, """flowchart LR
  U[Web/PWA] --> CF[Cloudflare Pages]
  U --> A[Supabase Auth]
  U --> P[PostgREST/RPC]
  P --> D[(PostgreSQL + RLS)]
  U --> E[parse-expense Edge Function]
  E --> D
  E --> G[Gemini API]""")
    add_bullets(doc, [
        "SAD chi tiết: docs/Solution_Architecture_Document_Family_Expense.docx.",
        "Nhãn UI Mục đích/Danh mục ánh xạ với bảng purposes/expense_types.",
        "Nhãn Tiền ra/Tiền vào vẫn tương thích enum database Chi tiêu/Thu nhập; dữ liệu Hoàn tiền/Tạm ứng cũ được bảo toàn.",
    ])

    heading(doc, "4. Codebase & Repositories", 1)
    add_table(doc, ["Hạng mục", "Vị trí", "Ghi chú"], [
        ("Frontend", "src/", "pages, components, context, lib, tests"),
        ("Database", "supabase/migrations/", "Migration additive; không sửa migration đã áp dụng"),
        ("Edge Functions", "supabase/functions/parse-expense/", "Deno + Supabase JS + Zod"),
        ("Import tooling", "scripts/import-excel.mjs", "Dry-run và migration dữ liệu Excel"),
        ("Tests", "src/**/*.test.*, tests/e2e", "Vitest/RTL và Playwright"),
        ("Config", "package.json, vite.config.ts, supabase/config.toml", "Không commit .env"),
        ("Docs", "README.md, CHANGELOG.md, docs/", "Nguồn tham chiếu và handoff"),
    ], [1.25, 2.4, 2.85])
    add_table(doc, ["Repository", "Remote", "Branch strategy", "Trạng thái"], [
        ("Family Expense", "TBD - workspace hiện chưa xác nhận Git remote", "Đề xuất: protected main + feature branches + PR", "Cần theo dõi"),
    ], [1.3, 2.5, 1.8, 0.9])

    heading(doc, "5. Environment & Infrastructure", 1)
    add_table(doc, ["Môi trường", "URL/Provider", "Backend", "Quyền truy cập", "Trạng thái"], [
        ("Local", "http://127.0.0.1:5173 / Vite", "Supabase local hoặc configured project", "Developer machine", "Có"),
        ("Preview", "Cloudflare Pages deployment URL", "Không dùng production mutation để test", "Cloudflare team", "Có theo deploy thủ công"),
        ("Staging", "TBD", "Đề xuất Supabase staging riêng", "TBD", "Chưa có"),
        ("Production", "family-expense-8fo.pages.dev", "Supabase Cloud linked project", "Cloudflare/Supabase admins", "Đang chạy"),
    ], [0.85, 1.75, 1.75, 1.25, 0.9])
    checklist(doc, [
        (False, "Ghi project ID/tenant name vào password manager hoặc CMDB, không ghi key trong tài liệu."),
        (False, "Xác nhận Site URL và Redirect URLs cho local/staging/production."),
        (False, "Bàn giao quyền owner tối thiểu 2 người cho Cloudflare và Supabase."),
    ])

    heading(doc, "6. Deployment & CI/CD", 1)
    heading(doc, "6.1 Quy trình hiện tại", 2)
    add_numbered(doc, [
        "Chạy pnpm lint, pnpm typecheck, pnpm test và pnpm build.",
        "Nếu có migration: supabase db push sau khi review và sao lưu phù hợp.",
        "Nếu đổi AI: supabase functions deploy parse-expense.",
        "Deploy thư mục dist lên Cloudflare Pages; kiểm tra HTTP 200 và smoke test auth/giao dịch.",
        "Cập nhật CHANGELOG với URL deployment và kết quả test.",
    ])
    add_callout(doc, "Trạng thái", "Quy trình production hiện thiên về thủ công. Cần chuyển sang Git-based CI/CD có approval và staging.", fill="FFF7E6", color="8A4B08")
    heading(doc, "6.2 Rollback", 2)
    add_numbered(doc, [
        "Frontend: chọn deployment Cloudflare Pages ổn định gần nhất và rollback/promote theo dashboard/CLI.",
        "Edge Function: redeploy artifact/commit trước đã kiểm thử.",
        "Database: không sửa/xóa migration cũ; ưu tiên forward-fix migration. Với sự cố dữ liệu, kích hoạt restore theo chính sách backup Supabase.",
        "Sau rollback: smoke test đăng nhập, tải danh sách, tạo/sửa giao dịch và AI; ghi incident timeline.",
    ])

    heading(doc, "7. Configuration & Secrets Management", 1)
    add_table(doc, ["Biến/credential", "Nơi sử dụng", "Nơi lưu", "Rotation"], [
        ("VITE_SUPABASE_URL", "Frontend", "Cloudflare env / .env.local", "Khi đổi project"),
        ("VITE_SUPABASE_ANON_KEY", "Frontend", "Cloudflare env / .env.local", "Theo thay đổi Supabase key"),
        ("GEMINI_API_KEY", "Edge Function", "Supabase Secret", "90 ngày hoặc khi nghi lộ"),
        ("GEMINI_MODEL", "Edge Function", "Supabase Secret/config", "Khi model đổi/deprecate"),
        ("SUPABASE_URL/ANON_KEY", "Edge Function", "Supabase Secret", "Khi đổi project/key"),
        ("Cloudflare API token", "Deploy/CI", "CI environment secret", "90 ngày; scope Pages project"),
        ("Supabase access token", "Migration/deploy", "CI environment secret", "90 ngày; scope tối thiểu"),
    ], [1.55, 1.35, 1.95, 1.65])
    add_numbered(doc, [
        "Tạo credential mới và cập nhật staging trước.",
        "Smoke test; cập nhật production trong maintenance window nếu cần.",
        "Thu hồi credential cũ sau khi xác minh không còn request sử dụng.",
        "Ghi ngày rotation, owner và phạm vi - không ghi giá trị secret.",
    ])

    heading(doc, "8. Third-party Integrations & Dependencies", 1)
    add_table(doc, ["Dịch vụ", "Mục đích", "Plan/license", "Owner", "Theo dõi"], [
        ("Supabase Cloud", "Auth, DB, RLS, RPC, Edge Function", "Free/hiện hành - cần xác minh", "TBD", "Quota, backup, egress, email rate limit"),
        ("Cloudflare Pages", "Hosting/CDN/PWA", "Free/hiện hành - cần xác minh", "TBD", "Build quota, domain, API token"),
        ("Google AI Studio/Gemini", "Structured AI suggestion", "Free Tier/hiện hành - cần xác minh", "TBD", "Quota, data policy, model lifecycle"),
        ("npm ecosystem", "Runtime/build dependencies", "Open source theo package", "Tech Lead", "License/SCA/deprecation"),
    ], [1.3, 1.75, 1.55, 0.8, 1.1])

    heading(doc, "9. Database & Data", 1)
    add_table(doc, ["Nhóm", "Đối tượng chính", "Ghi chú"], [
        ("Tenant/Auth", "families, family_members", "family_id + membership active + owner/member"),
        ("Phân loại", "purposes, expense_types, payment_methods, accounts, events, beneficiaries", "Không hard delete danh mục đã dùng"),
        ("Giao dịch", "transactions", "amount > 0, soft delete deleted_at, source/audit AI"),
        ("Kế hoạch", "budgets, recurring_transactions", "Schema có sẵn; UI nâng cao ngoài MVP"),
        ("AI audit", "ai_usage_logs", "Metadata tối thiểu, không API key/token"),
        ("Import audit", "import_batches và RPC liên quan", "Atomic batch, duplicate guard"),
    ], [1.25, 2.55, 2.70])
    heading(doc, "9.1 Backup/restore và retention", 2)
    checklist(doc, [
        (False, "Xác minh backup/PITR theo Supabase plan và ghi retention chính thức."),
        (False, "Thực hiện restore drill sang project không-production ít nhất hàng quý."),
        (True, "Giao dịch xóa thông thường dùng deleted_at và có màn hình Đã xóa/khôi phục."),
        (False, "Chốt retention thùng rác và ai_usage_logs với Product/Privacy owner."),
        (False, "Export dữ liệu định kỳ ngoài nhà cung cấp nếu RPO yêu cầu."),
    ])

    heading(doc, "10. Monitoring, Logging & Alerting", 1)
    add_table(doc, ["Nguồn", "Theo dõi", "Alert/Owner", "Trạng thái"], [
        ("Cloudflare", "Deploy/build errors, availability, Web Vitals", "TBD", "Cần cấu hình"),
        ("Supabase", "Auth errors, DB CPU/connections, slow queries, Edge errors", "TBD", "Có log nền tảng; alert TBD"),
        ("ai_usage_logs", "Success/error, latency, quota/rate limit", "Google AI Owner TBD", "Có dữ liệu cơ bản"),
        ("Client", "JS errors, PWA offline/JWT refresh failures", "Frontend Owner TBD", "Chưa có error tracking chuẩn"),
        ("Synthetic", "Login page, health/smoke route", "Incident Contact TBD", "Chưa có"),
    ], [1.1, 2.55, 1.55, 1.30])
    add_bullets(doc, [
        "Không log token, secret, email đầy đủ, nội dung AI hoặc chi tiết giao dịch.",
        "Mức đề xuất: P1 mất truy cập/lộ dữ liệu; P2 auth/CRUD/AI lỗi diện rộng; P3 lỗi đơn lẻ/UI.",
        "Alert phải có owner, escalation path, quiet hours và runbook link.",
    ])

    heading(doc, "11. Known Issues, Technical Debt & Workarounds", 1)
    add_table(doc, ["Mục", "Ảnh hưởng", "Workaround hiện tại", "Ưu tiên"], [
        ("Thiếu Git/CI/CD/staging", "Khó audit/rollback, rủi ro deploy", "Chạy quality checks và deploy thủ công", "P0"),
        ("JWT issued-at-future trên PWA", "Request Supabase thất bại sau khi app ngủ", "Refresh session khi foreground; đăng nhập lại nếu còn lỗi", "P1"),
        ("Email rate limit Supabase", "Đăng ký/reset bị giới hạn", "Chờ quota, tránh gửi lặp; cấu hình SMTP khi cần", "P1"),
        ("Free Tier Gemini", "429/quota/model thay đổi", "Không retry vô hạn; hiển thị lỗi; model qua env", "P1"),
        ("Enum DB cũ", "Tên nghiệp vụ UI và DB khác nhau", "Mapping Tiền ra/Tiền vào; giữ dữ liệu legacy", "P2"),
        ("Storage chưa dùng", "Chưa có chứng từ", "Ngoài MVP", "P3"),
        ("Monitoring chưa đầy đủ", "Phát hiện sự cố chậm", "Kiểm tra platform logs thủ công", "P1"),
    ], [1.8, 1.55, 2.35, 0.8])

    heading(doc, "12. Runbooks / SOPs", 1)
    heading(doc, "12.1 Người dùng không đăng nhập/reset được", 2)
    add_numbered(doc, [
        "Kiểm tra status Supabase Auth và redirect URL của domain hiện tại.",
        "Xác định lỗi invalid credentials, email rate limit hay JWT/clock skew; không yêu cầu user gửi token.",
        "Nếu rate limit: dừng gửi lặp, chờ cửa sổ quota hoặc kiểm tra SMTP custom.",
        "Nếu link reset về login: kiểm tra route /dat-lai-mat-khau và Site/Redirect URLs.",
        "Ghi sự cố đã ẩn email/PII và xác minh đăng nhập lại.",
    ])
    heading(doc, "12.2 JWT issued-at-future trên mobile/PWA", 2)
    add_numbered(doc, [
        "Yêu cầu bật ngày/giờ tự động trên thiết bị và kiểm tra chênh lệch thời gian.",
        "Đưa app ra foreground để session refresh; thử lại một lần có giới hạn.",
        "Nếu còn lỗi: đăng xuất/đăng nhập lại; kiểm tra service worker/version mới nhất.",
        "Đối chiếu Supabase Auth logs theo timestamp, không log JWT.",
    ])
    heading(doc, "12.3 AI trả 429/JSON không hợp lệ", 2)
    add_numbered(doc, [
        "Kiểm tra ai_usage_logs metadata và Supabase Edge Function logs.",
        "429: không retry liên tục; thông báo user nhập thủ công và kiểm tra quota/model.",
        "JSON/unknown category: xác minh GEMINI_MODEL, schema và danh mục family; giữ server validation.",
        "Không giảm kiểm soát bằng cách gọi Gemini trực tiếp từ frontend.",
    ])
    heading(doc, "12.4 Lỗi deploy", 2)
    add_numbered(doc, [
        "Dừng production rollout nếu lint/typecheck/test/build không đạt.",
        "Kiểm tra env names, SPA fallback, auth redirect và project target.",
        "Rollback Cloudflare deployment gần nhất; forward-fix database nếu migration đã chạy.",
        "Smoke test và ghi CHANGELOG/incident record.",
    ])
    heading(doc, "12.5 Nghi ngờ truy cập chéo family", 2)
    add_numbered(doc, [
        "Xếp P1 security; tạm dừng mutation/deploy liên quan nếu cần.",
        "Bảo toàn log metadata; không sao chép dữ liệu tài chính vào chat/ticket.",
        "Kiểm tra RLS policies, RPC SECURITY DEFINER và auth.uid/family_id.",
        "Khắc phục bằng migration mới, chạy negative tests và đánh giá phạm vi ảnh hưởng.",
        "Thông báo theo quy trình privacy/security đã phê duyệt.",
    ])

    heading(doc, "13. Testing & QA", 1)
    add_table(doc, ["Lớp test", "Lệnh", "Phạm vi", "Bắt buộc trước deploy"], [
        ("Lint", "pnpm lint", "ESLint/quy tắc mã", "Có"),
        ("Type", "pnpm typecheck", "TypeScript strict", "Có"),
        ("Unit/RTL", "pnpm test", "Domain, import, AI, pages/components", "Có"),
        ("Build", "pnpm build", "Production bundle + PWA", "Có"),
        ("E2E", "pnpm test:e2e", "Luồng chính trên browser", "Staging/release"),
        ("DB/RLS", "Supabase local/staging tests", "Cross-family, role, RPC, migration", "Cần bổ sung vào CI"),
    ], [1.0, 1.35, 2.55, 1.60])
    add_callout(doc, "Tài khoản test", "Không ghi credential trong tài liệu. Tạo user/family test riêng ở staging, lưu trong password manager của đội và không tái sử dụng production data.")

    heading(doc, "14. Documentation Index", 1)
    add_table(doc, ["Tài liệu", "Vị trí", "Trạng thái"], [
        ("Solution Architecture Document", "docs/Solution_Architecture_Document_Family_Expense.docx", "Đã hoàn thành"),
        ("Project Handoff Document", "docs/Handoff_Document_Family_Expense.docx", "Đã hoàn thành - còn TBD ký nhận"),
        ("README / Setup / Deploy", "README.md", "Đã có; rà soát khi pipeline đổi"),
        ("Change history", "CHANGELOG.md", "Đã cập nhật theo thay đổi"),
        ("Database schema/policies", "supabase/migrations/*.sql", "Nguồn sự thật kỹ thuật"),
        ("AI API contract", "supabase/functions/parse-expense/index.ts", "Nguồn sự thật kỹ thuật"),
        ("User guide", "README.md và hướng dẫn trong UI", "Cần tách thành tài liệu người dùng nếu bàn giao khách hàng"),
        ("API docs", "PostgREST/RPC schema + code", "Chưa có catalog API độc lập"),
    ], [2.05, 3.25, 1.20])

    heading(doc, "15. Outstanding Tasks & Roadmap", 1)
    add_table(doc, ["Ưu tiên", "Việc", "Owner", "Due", "Trạng thái"], [
        ("P0", "Khởi tạo Git remote, baseline tag, protected main", "TBD", "TBD", "Chưa làm"),
        ("P0", "CI lint/typecheck/test/build + Cloudflare preview", "TBD", "TBD", "Chưa làm"),
        ("P0", "Supabase/Cloudflare staging tách biệt", "TBD", "TBD", "Chưa làm"),
        ("P1", "RLS negative tests và migration rehearsal", "TBD", "TBD", "Chưa làm"),
        ("P1", "Monitoring/alert/error tracking không chứa PII", "TBD", "TBD", "Chưa làm"),
        ("P1", "Backup restore và rollback drill", "TBD", "TBD", "Chưa làm"),
        ("P1", "Chốt retention/log/privacy policy", "TBD", "TBD", "Cần quyết định"),
        ("P2", "Đánh giá migration enum Tiền vào/Tiền ra dài hạn", "TBD", "TBD", "Theo dõi"),
        ("P3", "Storage chứng từ/recurring/queue khi có business case", "TBD", "TBD", "Backlog"),
    ], [0.65, 2.95, 0.75, 0.65, 1.50])

    doc.add_page_break()
    heading(doc, "16. Sign-off", 1)
    add_table(doc, ["Bên", "Đại diện", "Vai trò", "Ngày", "Xác nhận/Ghi chú"], [
        ("Bàn giao", "TBD", "TBD", "__/__/____", "________________________"),
        ("Tiếp nhận", "TBD", "TBD", "__/__/____", "________________________"),
        ("Product Owner", "TBD", "TBD", "__/__/____", "________________________"),
        ("Security/Ops (nếu có)", "TBD", "TBD", "__/__/____", "________________________"),
    ], [1.20, 1.10, 1.15, 1.0, 2.05])
    checklist(doc, [
        (False, "Đã nhận quyền truy cập cần thiết và xác minh MFA."),
        (False, "Đã nhận tài liệu, mã nguồn, migration và lịch sử triển khai."),
        (False, "Đã thực hiện walkthrough kiến trúc, deploy, rollback và incident runbooks."),
        (False, "Đã xác nhận danh sách outstanding tasks, owner và thời hạn."),
        (False, "Đã thu hồi quyền không còn cần thiết của đội bàn giao."),
    ])

    for table in doc.tables:
        for row in table.rows:
            prevent_row_split(row)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_together = True

    doc.core_properties.title = "Project Handoff Document - Family Expense"
    doc.core_properties.subject = "Bàn giao dự án Family Expense"
    doc.core_properties.author = "Family Expense Delivery Team"
    doc.core_properties.keywords = "Family Expense, handoff, operations, Supabase, Cloudflare, Gemini"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
